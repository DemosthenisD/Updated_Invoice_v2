import streamlit as st
import pandas as pd
from docx import Document
import time
from io import BytesIO, StringIO
import base64
import os
from base64 import b64encode
import win32com.client
import pythoncom
from login import login

# Initialize invoices in session state
if 'invoices' not in st.session_state:
    st.session_state.invoices = []

# Custom CSS for the success message and animation
st.markdown("""
    <style>
    .hidden {
        display: none;
    }
    .success-message {
        font-size: 1.5rem;
        color: green;
        opacity: 0;
        transition: opacity 1s ease-in-out;
    }
    .success-message.show {
        opacity: 1;
    }
    </style>
    """, unsafe_allow_html=True)

# JavaScript to show the success message
st.markdown("""
    <script>
    function showSuccessMessage() {
        var successMessage = document.getElementById("successMessage");
        successMessage.classList.add("show");
    }
    </script>
    """, unsafe_allow_html=True)


def fill_placeholders(doc, data):
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text


    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in data.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, str(value))
                                        inline[i].text = text

def load_dataframe(file_path, worksheet_name):
    try:
        df = pd.read_excel(file_path, sheet_name= worksheet_name)
        return df
    except FileNotFoundError:
        st.write(f"File {file_path} not found.")
        exit()


def main():
    if not st.session_state.logged_in:
            login()
    else:
        # Load the template document
        template_path = 'template1.docx'
        template_doc = Document(template_path)

        file_path = os.path.join(os.getcwd(), 'InvoiceLogTemplate.xlsx')  # Full file path
        worksheet_name = "Clients"
        df = load_dataframe(file_path, worksheet_name)

        file_path_csv = 'new_record.csv'

        # Placeholder for storing records
        records = []

        home_page      = st.sidebar.page_link('App.py',                  label="HOME",           icon="🏡")
        record_page    = st.sidebar.page_link('pages/record.py',         label="RECORD",         icon="📓")    
        add_new_record = st.sidebar.page_link('pages/add_new_record.py', label="ADD NEW RECORD", icon="✒️")    
        
        clients  = df['Client'].unique()
        projects = df['Project'].unique()

        client = st.selectbox("Select Client", clients)
        # if client:
        invoice_no = len(st.session_state.invoices) + 1
        col1,col2 = st.columns([1,1])
        with col1:
            date = st.date_input("Date")
        with col2:
            amount = st.number_input("Amount")

        vat = st.select_slider('VAT', options=[i for i in range(0, 101)], format_func=lambda x: f'{x}%')


        filtered_client_code = df[df['Client'] == client]['client_code'].unique()

        filtered_projects = df[df['Client'] == client]['Project'].unique()          
        project = st.selectbox("Select Project", filtered_projects)
        description = st.text_area("Description")
        address = "My Address"
        vat_number = "My VAT No"
        year = date.year

        # BUTTONS
        col1, col2,col3 = st.columns([1,1,2])
        with col1:
            generate_invoice = st.button('Generate Invoice', key="generate")
        with col2:
            save_record_button      = st.button("Save Record", key="save_record")
        with col3:
            pass    

        # Save Record Button
        if save_record_button:
            add_new_record = {
                'Client' : client,
                'Project': project,
                'Date Issued': date,
                'Year': year,
                'client_code': filtered_client_code,
            }

            try:
                # Read existing data from the Excel file
                xl = pd.ExcelFile(file_path)
                
                # Load all sheets into a dictionary of DataFrames
                dfs = {sheet_name: xl.parse(sheet_name) for sheet_name in xl.sheet_names}
                
                # Update the specific sheet with the new record
                if 'InvoiceLogTemplate' in dfs:
                    df = dfs['InvoiceLogTemplate']
                    new_record_df = pd.DataFrame([add_new_record])
                    df = pd.concat([df, new_record_df], ignore_index=True)
                    dfs['InvoiceLogTemplate'] = df
                else:
                    # Handle case where 'InvoiceLogTemplate' sheet doesn't exist
                    dfs['InvoiceLogTemplate'] = pd.DataFrame([add_new_record])
                
                # Write all sheets back to the Excel file
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    for sheet_name, df in dfs.items():
                        df.to_excel(writer, index=False, sheet_name=sheet_name)
                
                st.success("Record saved successfully.")
            except PermissionError:
                st.error("PermissionError: Permission denied. Please make sure the file is not open elsewhere and you have write permission.")
            except Exception as e:
                st.error(f"An error occurred: {e}")


        if generate_invoice:
            # Define data to fill in placeholders
            vat_value = (amount * vat)/100
            total_invoice = amount + vat_value
            data = {
                '{{placeholder1}}': client,
                '{{placeholder2}}': address,
                '{{placeholder3}}': vat_number,
                '{{placeholder4}}': date,
                '{{placeholder5}}': invoice_no,
                '{{placeholder6}}': year,
                '{{placeholder7}}': description,
                '{{placeholder8}}': amount,
                '{{placeholder9}}': vat_value,
                '{{placeholder10}}':total_invoice
                # Add more placeholders as needed
            }
            # Fill placeholders
            fill_placeholders(template_doc, data)

            # Simulate invoice generation
            with st.spinner('Generating invoice...'):
                time.sleep(4)  # Simulate time taken to generate the invoice

                # Save invoice to session
                st.session_state.invoices.append({
                    'client': client,
                    'address': address,
                    'vat_number': vat_number,
                    'date': date,
                    'invoice_no': invoice_no,
                    'year': year,
                    'description': description,
                    'amount': amount,
                    'vat_value': vat_value,
                    'total_invoice': total_invoice,
                    'download_format': None  # Initialize download format
                })
                st.success('Invoice generated successfully!')

        # Display download section only if invoices are generated
        if st.session_state.invoices:
            try:
                invoice = st.session_state.invoices[-1]  # Get the last generated invoice
                st.markdown("### Download it:")
                # Select download format
                format_option = st.selectbox("Select download format", ["DOCX","PDF"], key="format_option")
                invoice['download_format'] = format_option  # Update download format in session
                if format_option == "DOCX":               
                    # Save the docx file in the root directory
                    tmp_download_link = download_link_docx(template_doc, 'filled_document.docx', 'Click here to download DOCX')
                    st.markdown(tmp_download_link, unsafe_allow_html=True)

                elif format_option == "PDF":
                    # Convert the document to PDF
                    pdf_file = convert_to_pdf('filled_document.docx')
                    tmp_download_link = download_link_pdf(pdf_file, 'filled_document.pdf', 'Click here to download PDF')
                    st.markdown(tmp_download_link, unsafe_allow_html=True)
                    remove_document_file('filled_document.docx')  # Adjust this path as per your actual file name
                    remove_document_file('filled_document.pdf')  # Adjust this path as per your actual file name
            except:
                    tmp_download_link = download_link_docx(template_doc, 'filled_document.docx', 'Click here to download DOCX')
                    st.markdown(tmp_download_link, unsafe_allow_html=True)


# Function to generate download link
def download_link_pdf(file_path, text, label):
    with open(file_path, 'rb') as f:
        data = f.read()
    href = f'<a href="data:application/octet-stream;base64,{b64encode(data).decode()}" download="{file_path}">{label}</a>'
    # Cleanup: Remove the docx file after PDF conversion and download
    return href

# Assuming your template_doc is a docx.Document object
def download_link_docx(doc, filename, text):
    """Generates a download link for a Docx file."""
    root_dir = os.getcwd()  # Get the current working directory (root directory)
    doc.save(os.path.join(root_dir, filename))  # Save the docx file to the root directory
    with open(os.path.join(root_dir, filename), 'rb') as f:
        doc_bytes = f.read()
    href = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_bytes).decode()}" download="{filename}">{text}</a>'
    return href

def convert_to_pdf(docx_file):
    # Initialize COM
    pythoncom.CoInitialize()
    try:
        # Get the absolute path of the DOCX file
        docx_path = os.path.abspath(docx_file)
        # Generate the PDF file name
        pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
        # Create an instance of the Word application
        word = win32com.client.Dispatch("Word.Application")
        try:
            # Open the DOCX file
            doc = word.Documents.Open(docx_path)
            # Save the document as PDF
            doc.SaveAs(pdf_path, FileFormat=17)  # 17 is the PDF file format
            doc.Close()
        except Exception as e:
            raise e
        finally:
            # Close the Word application
            word.Quit()
    finally:
        # Uninitialize COM
        pythoncom.CoUninitialize()
    return pdf_path

# Function to remove the downloaded document file from root directory
def remove_document_file(file_path):
    """Removes the document file from root directory."""
    if os.path.exists(file_path):
        os.remove(file_path)

if __name__ == "__main__":
    main()

































