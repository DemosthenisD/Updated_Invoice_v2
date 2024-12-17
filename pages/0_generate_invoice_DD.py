import streamlit as st
import pandas as pd
from docx import Document
import time
from io import BytesIO, StringIO
import base64
import os
from base64 import b64encode
from streamlit_extras.switch_page_button import switch_page
import convertapi
from streamlit_free_text_select import st_free_text_select


st.sidebar.page_link('pages/0_generate_invoice_DD.py',     label="Generate Invoice",               icon="🏡")
st.sidebar.page_link('pages/1_list_of_clients_projects.py',label="List of Clients / Projects List",icon="📓")    
st.sidebar.page_link('pages/2_add_new_client_project.py',  label="Add New Client/Project record",  icon="✒️")  


# Define SessionState class
class SessionState:
    def __init__(self, **kwargs):
        self._state = kwargs

    def __getattr__(self, item):
        return self._state.get(item, None)

    def __setattr__(self, key, value):
        if key == '_state':
            super().__setattr__(key, value)
        else:
            self._state[key] = value



# Create a SessionState object to store session variables
session_state = SessionState(invoices=[])

# Set your ConvertAPI secret key
convertapi.api_secret = 'WNkNerQr6LJX6JUw'  # Replace with your actual API key

# Initialize invoices in session state
if 'invoices' not in st.session_state:
    st.session_state.invoices = []




# Custom CSS for the success message and animation
st.markdown("""
    <style>
    .hidden {
        display: none;
    }
    .success-message {
        font-size: 1.5rem;
        color: green;
        opacity: 0;
        transition: opacity 1s ease-in-out;
    }
    .success-message.show {
        opacity: 1;
    }
    </style>
    """, unsafe_allow_html=True)

# JavaScript to show the success message
st.markdown("""
    <script>
    function showSuccessMessage() {
        var successMessage = document.getElementById("successMessage");
        successMessage.classList.add("show");
    }
    </script>
    """, unsafe_allow_html=True)


def fill_placeholders(doc, data):
    for p in doc.paragraphs:
        for key, value in data.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if key in inline[i].text:
                        text = inline[i].text.replace(key, str(value))
                        inline[i].text = text


    for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in data.items():
                            if key in p.text:
                                inline = p.runs
                                for i in range(len(inline)):
                                    if key in inline[i].text:
                                        text = inline[i].text.replace(key, str(value))
                                        inline[i].text = text

def load_dataframe(file_path, worksheet_name):
    try:
        df = pd.read_excel(file_path, sheet_name= worksheet_name)
        return df
    except FileNotFoundError:
        st.write(f"File {file_path} not found.")
        exit()


# Function to generate download link
def download_link_pdf(file_path, text, label):
    with open(file_path, 'rb') as f:
        data = f.read()
    href = f'<a href="data:application/octet-stream;base64,{b64encode(data).decode()}" download="{file_path}">{label}</a>'
    # Cleanup: Remove the docx file after PDF conversion and download
    return href

# Assuming your template_doc is a docx.Document object
def download_link_docx(doc, year, invoice_no, client, filename, text):
    """Generates a download link for a Docx file."""
    root_dir = os.getcwd()  # Get the current working directory (root directory)
    save_file_name = f"{year}/{invoice_no} {client} Invoice"
    doc.save(os.path.join(root_dir , filename))  # Save the docx file to the root directory
    with open(os.path.join(root_dir, filename), 'rb') as f:
        doc_bytes = f.read()
    href = f'<a href="data:application/octet-stream;base64,{base64.b64encode(doc_bytes).decode()}" download="{save_file_name}.docx">{text}</a>'
    return href



def convert_docx_to_pdf(file_path):
    # Use ConvertAPI to convert the DOCX file to PDF
    result = convertapi.convert('pdf', {
        'File': file_path
    }, from_format='docx')

    # Save the resulting PDF to a temporary file
    pdf_file_path = 'converted.pdf'
    result.save_files(pdf_file_path)
    
    # Read the PDF file and return its content
    with open(pdf_file_path, 'rb') as f:
        pdf_file = f.read()

    # Clean up temporary files
    os.remove(pdf_file_path)

    return pdf_file

# def convert_to_pdf(docx_file):
#     # Initialize COM
#     pythoncom.CoInitialize()
#     try:
#         # Get the absolute path of the DOCX file
#         docx_path = os.path.abspath(docx_file)
#         # Generate the PDF file name
#         pdf_path = os.path.splitext(docx_path)[0] + ".pdf"
#         # Create an instance of the Word application
#         word = win32com.client.Dispatch("Word.Application")
#         try:
#             # Open the DOCX file
#             doc = word.Documents.Open(docx_path)
#             # Save the document as PDF
#             doc.SaveAs(pdf_path, FileFormat=17)  # 17 is the PDF file format
#             doc.Close()
#         except Exception as e:
#             raise e
#         finally:
#             # Close the Word application
#             word.Quit()
#     finally:
#         # Uninitialize COM
#         pythoncom.CoUninitialize()
#     return pdf_path



# Function to remove the downloaded document file from root directory
def remove_document_file(file_path):
    """Removes the document file from root directory."""
    if os.path.exists(file_path):
        os.remove(file_path)


def main():
    if 'username' in st.session_state:

        file_path = os.path.join(os.getcwd(), 'InvoiceLogTemplate_DD_28062024.xlsx')  # Full file path - DD_04062024: UPDATED FILE NAME
        
        worksheet_project_list = "Project_List"  # DD_04062024: previously "Clients"
        df_project_list = load_dataframe(file_path, worksheet_project_list)

        worksheet_client_list = "Client_List" 
        df_client_list = load_dataframe(file_path, worksheet_client_list)
        
        clients  = df_project_list['Client'].unique()

        df_Invoice_List = load_dataframe(file_path,"InvoiceLogTemplate")

        col1, col2, col3 = st.columns([1,1,1])
        with col1:
            client = st.selectbox("Select Client",clients)
             # DD_28062024: Added next row
            Client_Name_For_Invoice = df_project_list[df_project_list['Client'] == client]['Client Name (for Invoices)'].unique()
        with col2:
            filtered_address = df_client_list[df_client_list['Client'] == client]['Address'].unique() 
            address = st_free_text_select(label="Address", 
                                          options=filtered_address,
                                            format_func=lambda x: x.lower(),
                                            placeholder="Select or Type Address",
                                            disabled=False,
                                            delay=300,)
        with col3:
            vat_number = df_project_list[df_project_list['Client'] == client]['VAT_No'].dropna().unique().tolist()  # DD_17122024 added the dropna & "tolist"# DD_04062024: previously "My VAT No"
            # DD_17122024 Add default VAT options (0% and 19%) if not already present
            vat_number = vat_number if len(vat_number) > 0 else ["No VAT No Found"] # DD_17122024 using default if empty list
            # DD_17122024 removed after code worked: st.write(f"VAT No: {vat_number}")
            # # DD_28062024: previously--> vat_no     = st.selectbox("VAT No", vat_number)
            vat_no = st_free_text_select(label="VAT No", 
                                        options=vat_number, #format_func=lambda x: x.lower(),
                                        placeholder="Select or Type VAT Number",
                                        disabled=False,
                                        delay=300,)
        # if client:
       
        col1, col2, col3 = st.columns([1,1,1]) # DD_28062024: previously--> st.columns([1,1])
        with col2:            # DD_28062024: previously--> with col1
            date = st.date_input("Date")
            year = date.year
        with col3:            # DD_28062024: previously--> with col2
            amount = st.number_input("Amount",step=1,value=0)  # DD_17122024 Default value to prevent None
        with col1:            # DD_28062024: New addition
            Pre_invoice_no = len(df_Invoice_List[df_Invoice_List['Year'] == year]) +1
            #invoice_no= year &"/"& st.text_input(label="invoice No", value=Pre_invoice_no) # DD_28062024: previously--> _invoice_no = len(st.session_state.invoices) + 1
            invoice_no= st.text_input(label="invoice No", value=Pre_invoice_no)
            
        filtered_vat_pct = df_project_list[df_project_list['Client'] == client]['VAT %'].dropna().unique().tolist()  # DD_17122024 added the dropna & "tolist"
        default_vat_options = ['0%', '19%']
        filtered_vat_pct = filtered_vat_pct if len(filtered_vat_pct) > 0 else default_vat_options # DD_17122024 added the "tolist" and using default if empty list
        
        # DD_28062024: previously--> vat = st.selectbox("VAT %", filtered_vat,)                                           
        vat   = st_free_text_select(label="VAT %", 
                                    options=filtered_vat_pct,
                                    #format_func=lambda x: x.lower(),
                                    placeholder="Select or Type VAT %-age",
                                    disabled=False,
                                    delay=300,)
        vat = vat or 0        # DD_17122024 Defaults to 0 if vat is None also edited to make string into float
        vat = float(vat)
        #vat=convert_to_number(vat)
        
        filtered_client_code = df_project_list[df_project_list['Client'] == client]['client_code'].unique()

        filtered_projects = df_project_list[df_project_list['Client'] == client]['Project'].unique()          
        project = st_free_text_select(label="Select Project", 
                                      options=filtered_projects,
                                      format_func=lambda x: x.lower(),
                                      placeholder="Select or Enter a project",
                                      disabled=False,
                                      delay=300,)

        filtered_description = df_project_list[df_project_list['Client'] == client]['description'].unique() 
        description = st_free_text_select(label="Description",
                                          options=filtered_description,
                                            format_func=lambda x: x.lower(),
                                            placeholder="Select or Enter a description",
                                            disabled=False,
                                            delay=300,)

        
        with st.expander("Select Invoice Template and Format"):
            col1, col2 = st.columns([1,1])
            with col1:
                Filtered_Invoice_Template = (df_project_list[df_project_list['Client'] == client]['Invoice Template'].unique()) # DD_28062024: previously-->  options_for_templates = (df_project_list[df_project_list['Client'] == client]['Invoice Template'].unique())
                # DD_28062024: ADDED NEXT 2 ROWS
                options_for_templates =df_project_list['Invoice Template'].unique().tolist()
                index_Selected = options_for_templates.index(Filtered_Invoice_Template)
                # DD_28062024: added index in st.radio below
                invoice_template = st.radio("Select Template for Invoice", options_for_templates, index=index_Selected, key="invoice_template")
            with col2:
                # Select download format
                format_option = st.radio("Select download format", ["DOCX", "PDF"],  key="format_option")
        # DD 17122024 Added debugging statements to verify the values of amount and vat just before this calculation
        st.write(f"Amount: {amount}, VAT: {vat}")
        VAT_Amount = (amount * vat)/100 + 0
        Expenses_Net_Amount = 0
        Expenses_VAT_Amount = 0 #(Expenses_Net_Amount*vat)/100
        
        # BUTTONS
        col1, col2,col3 = st.columns([1,1,2])
        with col1:
            generate_invoice = st.button('Generate Invoice', key="generate")
        with col2:
            save_record_button      = st.button("Save Record", key="save_record")
        with col3:
            pass    
    
        # Save Record Button
        if save_record_button:
            add_new_record = {
                'Year': year,
                'Client' : client,
                'Legal Entity': Client_Name_For_Invoice,
                'Project': project,
                'client_code': filtered_client_code,
                'Type': "Invoice",
                'Invoice': None, # DD_04062024 added this
                'Expenses_Net_Amt': 0, #Expenses_Net_Amount, # DD_28062024 added this
                #'Expenses_VAT_Amt': 0, #Expenses_VAT_Amount, # DD_28062024 added this 
                'Invoice No': invoice_no,       # DD_04062024 added this
                'Invoiced Amt Net': amount, # DD_04062024 added this
                'VAT_Amount': VAT_Amount,       # DD_28062024 edited this
                'Date Issued': date,
                'Address': address,
                'description': description,
                'VAT_No': vat_no,         
            }

            try:
                # Read existing data from the Excel file
                xl = pd.ExcelFile(file_path)
                
                # Load all sheets into a dictionary of DataFrames
                dfs = {sheet_name: xl.parse(sheet_name) for sheet_name in xl.sheet_names}
                
                # Update the specific sheet with the new record
                if 'InvoiceLogTemplate' in dfs:
                    df = dfs['InvoiceLogTemplate']
                    new_record_df = pd.DataFrame([add_new_record])
                    df = pd.concat([df, new_record_df], ignore_index=True)
                    dfs['InvoiceLogTemplate'] = df
                else:
                    # Handle case where 'InvoiceLogTemplate' sheet doesn't exist
                    dfs['InvoiceLogTemplate'] = pd.DataFrame([add_new_record])
            
                # Write all sheets back to the Excel file
                with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
                    for sheet_name, df in dfs.items():
                        df.to_excel(writer, index=False, sheet_name=sheet_name)

                st.success("Record Saved Successfully.")
            except PermissionError:
                st.error("PermissionError: Permission denied. Please make sure the file is not open elsewhere and you have write permission.")
            except Exception as e:
                st.error(f"An error occurred: {e}")

        # Check if download format is saved in session
        if 'download_format' not in st.session_state:
            st.session_state.download_format = 'DOCX'
        
        if generate_invoice:
            try:
                # Define data to fill in placeholders
                vat_value = VAT_Amount
                total_invoice = amount + vat_value + Expenses_Net_Amount + Expenses_VAT_Amount
                data = {
                    '{{placeholder1}}': Client_Name_For_Invoice,
                    '{{placeholder2}}': address,
                    '{{placeholder3}}': vat_no,
                    '{{placeholder4}}': date,
                    '{{placeholder5}}': f"{year}-{invoice_no}",
                    #'{{placeholder6}}': year,  # DD_28062024 REMOVED THIS since invoice_no now inlcudes year
                    '{{placeholder7}}': description,
                    '{{placeholder8}}': amount,
                    '{{placeholder8_Exp}}': Expenses_Net_Amount,
                    '{{placeholder8_Tot}}': amount+Expenses_Net_Amount,
                    '{{placeholder9}}': vat_value,
                    '{{placeholder9_Exp}}': Expenses_VAT_Amount,
                    '{{placeholder9_Tot}}': vat_value + Expenses_VAT_Amount,
                    '{{placeholder10}}':total_invoice,
                    # Add more placeholders as needed
                }

                # Save invoice to session
                st.session_state.invoices.append({
                    # DD_28062024 Added following 16 rows and commented previous versions
                    'Year': year,
                    'Client' : client,
                    'Legal Entity': Client_Name_For_Invoice,
                    'Project': project,
                    'client_code': filtered_client_code,
                    'Type': "Invoice",
                    'Invoice': None, # DD_04062024 added this
                    #'Expenses_Net_Amt': Expenses_Net_Amount # DD_28062024 added this
                    #'Expenses_VAT_Amt': Expenses_VAT_Amount  # DD_28062024 added this 
                    'Invoice No': invoice_no,       # DD_04062024 added this
                    'Invoiced Amt Net': amount, # DD_04062024 added this
                    'VAT_Amount': VAT_Amount,       # DD_28062024 edited this
                    'Date Issued': date,
                    'Address': address,
                    'description': description,
                    'VAT_No': vat_no,
                    #'client': client,
                    #'address': address,
                    #'vat_number': vat_number,
                    #'date': date,
                    #'invoice_no': invoice_no,
                    #'year': year,
                    #'description': description,
                    #'amount': amount,
                    #'vat_value': vat_value,
                    #'total_invoice': total_invoice,
                    'invoice_template': invoice_template,  # Initialize Invoice Template
                    'download_format' : None  # Initialize download format
                })
                if invoice_template == "Template-1":
                    template_path = 'template1_v3.docx' # DD_28062024 Edited this
                    template_doc = Document(template_path)

                    # Fill placeholders
                    fill_placeholders(template_doc, data)
                    # Simulate invoice generation
                    with st.spinner('Generating invoice...'):
                        time.sleep(4)  # Simulate time taken to generate the invoice
                        
                    download_section(template_doc, year, invoice_no, client, format_option)
                    

                elif invoice_template == "Template-2":
                    # st.error("Template 2 does not exists")
                    template_path = 'template2_v3.docx' # DD_28062024 Edited this
                    template_doc = Document(template_path)

                    # Fill placeholders
                    fill_placeholders(template_doc, data)
                    # Simulate invoice generation
                    with st.spinner('Generating invoice...'):
                        time.sleep(4)  # Simulate time taken to generate the invoice
                        
                    download_section(template_doc, year, invoice_no, client, format_option)



                # Store the template and other information in session
                session_state.template_doc = template_doc
                session_state.invoice_no = invoice_no
                session_state.client = client
                session_state.invoice_generated = True  # Mark invoice as generated

            except Exception as e:
                pass
                # st.warning("Select Invoice Template")

    else:
        st.error("There's some issue, Its requires to login again your app!")
        if st.button("Verify Auth"):
            switch_page('app')


def download_section(template_doc, year, invoice_no, client, format_option):
            # Display download section only if invoices are generated
            if st.session_state.invoices:
                try:
                    invoice = st.session_state.invoices[-1]  # Get the last generated invoice
                    
                    invoice['download_format'] = format_option  # Update download format in session
                    if format_option == "DOCX":
                        st.markdown("### Download it:")       
                        # Save the docx file in the root directory
                        tmp_download_link = download_link_docx(template_doc, year, invoice_no, client, "filled_document.docx", 'Click here to download DOCX')
                        st.markdown(tmp_download_link, unsafe_allow_html=True)
                        st.success('Invoice generated successfully!')
                        download_link_docx(template_doc, year, invoice_no, client, "filled_document.docx")

                    elif format_option == "PDF":
                        filename = "filled_document.docx"
                        root_dir = os.getcwd()  # Get the current working directory (root directory)
                        save_file_name = f"{year}-{invoice_no} {client} Invoice"
                        template_doc.save(os.path.join(root_dir , filename))  # Save the docx file to the root directory
  
                        docx_file_path = 'filled_document.docx'

                        if os.path.exists(docx_file_path):
                            try:
                                pdf_file = convert_docx_to_pdf(docx_file_path)

                                if pdf_file:
                                    st.markdown("### Download it:")
                                    st.success("Invoice generated successfully!")
                                    
                                    # Create a download link
                                    b64 = base64.b64encode(pdf_file).decode()
                                    label = 'Click here to download PDF'
                                    href = f'<a href="data:application/pdf;base64,{b64}" download="{save_file_name}.pdf">{label}</a>'                              
                                    st.markdown(href, unsafe_allow_html=True)
                                else:
                                    st.error("Failed to convert the DOCX file to PDF.")
                            except Exception as e:
                                st.error(f"An error occurred during conversion: {e}")
                        else:
                            st.error(f"File '{docx_file_path}' not found in the root directory.")   

                        # remove_document_file('filled_document.docx')  # Adjust this path as per your actual file name
                        # remove_document_file('filled_document.pdf')  # Adjust this path as per your actual file name
                except:
                        file_name = f"{invoice_no}-{client}.docx"
                        tmp_download_link = download_link_docx(template_doc, file_name, 'Click here to download DOCX')
                        st.markdown(tmp_download_link, unsafe_allow_html=True)    


def convert_to_number(input_text):
    try:
        # Try to convert to an integer
        number = int(input_text)
        return number
    except ValueError:
        try:
            # If conversion to integer fails, try to convert to float
            number = float(input_text)
            return number
        except ValueError:
            # If conversion to both integer and float fails, return None or raise an error
            return None
            # Alternatively, you can raise an error
            # raise ValueError(f"Cannot convert '{input_text}' to a number")




if __name__ == "__main__":
    main()

































